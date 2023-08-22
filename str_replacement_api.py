# -*- coding: utf-8 -*-
###########################################################################
#    Copyright 2023 奇客罗方智能科技 https://www.geekercloud.com
#    ITGeeker.net <alanljj@gmail.com>
############################################################################
import glob
import os
from docx import Document


def do_replace_old_str_action(old_str, new_str, prgrp):
    if old_str in prgrp.text:
        prgrp.text = prgrp.text.replace(old_str, new_str)
        inline = prgrp.runs
        for i in range(len(inline)):
            if old_str in inline[i].text:
                text = inline[i].text.replace(old_str, new_str)
                inline[i].text = text


def docx_replace_old_to_new(doc_obj, old_string, new_string):
    try:
        # for paragraph in doc_obj.paragraphs:
        #     for run in paragraph.runs:
        #         if old_string in run.text:
        #             run.text = run.text.replace(old_string, new_string)

        for p in doc_obj.paragraphs:
            print('p.text: ', p.text)
            if old_string in p.text:
                p.text = p.text.replace(old_string, new_string)
                inline = p.runs
                # Loop added to work with runs (strings with same style)
                for i in range(len(inline)):
                    # print('inline[i].text: ', inline[i].text)
                    if old_string in inline[i].text:
                        text = inline[i].text.replace(old_string, new_string)
                        # text = old_string.sub(new_string, inline[i].text)
                        # text = regex.sub(replace, inline[i].text, flags=re.IGNORECASE)
                        inline[i].text = text
        for table in doc_obj.tables:
            for row in table.rows:
                for cell in row.cells:
                    # docx_replace_old_to_new(cell, old_string, new_string)
                    for paragraph in cell.paragraphs:
                        if old_string in paragraph.text:
                            print('paragraph.text: ', paragraph.text)
                            paragraph.text = paragraph.text.replace(old_string, new_string)
        for section in doc_obj.sections:
            for ph in section.header.paragraphs:
                do_replace_old_str_action(old_string, new_string, ph)
            for pf in section.footer.paragraphs:
                do_replace_old_str_action(old_string, new_string, pf)
        return True
    except Exception as err:
        print('err@replace old to new: %s' % err)
        return False


def generate_file_and_str_list(report_p, val_list):
    replaced_p = os.path.join(report_p, '已替换文件')
    print('replaced_p: ', replaced_p)
    if not os.path.exists(replaced_p):
        os.makedirs(replaced_p)
    docx_list = [f for f in glob.glob(report_p + r"\[!~$]*.docx")]
    print('docx_list: ', docx_list)
    replaced = replace_str_for_file_list(docx_list, val_list, replaced_p)
    if replaced:
        return len(docx_list)
    return False


def replace_str_for_file_list(docx_list, val_list, replaced_p):
    for docx_f in docx_list:
        doc = Document(docx_f)
        filename, file_extension = os.path.splitext(docx_f)
        basename_no_ext = os.path.splitext(os.path.basename(docx_f))[0]
        # print('filename: ', filename)
        print('file_extension: ', file_extension)
        print('basename: ', basename_no_ext)

        replaced_file = os.path.join(replaced_p, basename_no_ext + '-replaced' + file_extension)
        print('replaced_file: ', replaced_file)

        for val in val_list:
            old_str = val.get('org_str')
            new_str = val.get('replaced_str')
            rplc_new = docx_replace_old_to_new(doc, str(old_str), str(new_str))
            if not rplc_new:
                return False
        doc.save(replaced_file)
    return True
